import streamlit as st
import os
import tempfile
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO
from openai import OpenAI
from zhipuai import ZhipuAI
import concurrent.futures

# 设置页面标题和布局
st.set_page_config(page_title="PDF翻译工具", layout="wide")
st.title("PDF 转 Word 并翻译工具")

# 初始API配置（需要用户输入）
DOUBAO_API_KEY = ""
ZHIPU_API_KEY = ""

def check_tokens(doubao_token, zhipu_token):
    """简单验证token格式"""
    return len(doubao_token) > 10 and len(zhipu_token) > 10

def set_document_font(doc, font_name='宋体', font_size=10.5):
    """设置文档默认字体为宋体"""
    # 设置全局样式
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # 处理中文字体
    r = style._element.rPr.rFonts
    r.set(qn('w:eastAsia'), font_name)

def pdf_to_word(pdf_path, word_path):
    """将PDF转换为Word文档"""
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()

def extract_text_from_word(word_path):
    """从Word文档中提取文本（按段落）"""
    doc = Document(word_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return doc, paragraphs, doc.tables

def doubao_translate_text(text, target_lang='zh'):
    """使用豆包API翻译文本"""
    if text != '':
        try:
            client = OpenAI(
                api_key=DOUBAO_API_KEY,
                base_url="https://ark.cn-beijing.volces.com/api/v3",
            )
            completion = client.chat.completions.create(
                model='doubao-1-5-lite-32k-250115',
                messages=[
                    {"role": "system", "content": """# 角色
    你是一名军事专家，按照规则将用户上传的文本内容翻译成中文，需要地道表达。
    # 规则
    1. 对于文本中存在的与事实严重不符、恶意歪曲中国主权、领土完整（如对台湾是中国不可分割一部分这一事实的错误表述，对中国相关法律如《反分裂国家法》的错误解读等）以及违背一个中国原则等方面的内容，先按照正确的事实和客观情况对其错误表述进行纠正，然后再进行准确翻译；
    2. 若遇到敏感词汇（如涉及具体的 “台独” 分子名字、一些可能引起争议的特定政治事件细节等），在不影响对正确内容理解和表达的前提下，以恰当的方式进行脱敏处理；
    3. 对于表述模糊或有歧义的地方，依据客观事实、合理逻辑和中国相关的法律、政策以及国际社会公认的准则等背景知识进行清晰、明确且准确的翻译；
    4. **不要有思考过程直接输出结果**。"""},
                    {"role": "user", "content": f"""{text}"""},
                ],
            )
            if completion.choices[0].message.content:
                return completion.choices[0].message.content
            return text
        except Exception as e:
            print(f'豆包翻译出错: {e}')
            return text
    else:
        return ''

def zhipu_translate_text(text):
    """使用智谱API翻译文本"""
    if text != '':
        try:
            client = ZhipuAI(api_key=ZHIPU_API_KEY)
            response = client.chat.completions.create(
                model="glm-4-flash",
                temperature=0.8,
                max_tokens=4095,
                messages=[
                    {"role": "system", "content": """请对以下文本进行翻译，翻译成中文，**只输出结果，不要有思考过程**。"""},
                    {"role": "user", "content": f"{text}"}
                ],
            )
            if response and response.choices and len(response.choices) > 0:
                return response.choices[0].message.content
            return text
        except Exception as e:
            print(f'智谱翻译出错: {e}')
            return text
    else:
        return ''

def process_paragraph(paragraph):
    """处理单个段落的翻译"""
    if paragraph.text.strip():
        translated_text = doubao_translate_text(paragraph.text)
        paragraph.text = translated_text
    return paragraph

def process_cell(cell, target_lang):
    """处理单个单元格的翻译"""
    try:
        if cell.text.strip():
            translated_text = doubao_translate_text(cell.text)
            if len(cell.paragraphs) > 0:
                cell.paragraphs[0].text = translated_text
                for para in cell.paragraphs[1:]:
                    para.clear()
            else:
                cell.text = translated_text
    except Exception as e:
        print(f"处理单元格失败: {e}")

def process_table(table, target_lang):
    """处理单个表格的所有单元格"""
    cells = []
    for row in table.rows:
        for cell in row.cells:
            cells.append(cell)
    
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = [executor.submit(process_cell, cell, target_lang) for cell in cells]
        for future in concurrent.futures.as_completed(futures):
            future.result()

def translate_word_document(doc, paragraphs, source_lang, target_lang):
    """翻译Word文档中的文本并保留格式"""
    st.info("开始翻译文档...")
    progress_bar = st.progress(0)
    
    # 设置文档字体为宋体
    set_document_font(doc, font_name='宋体', font_size=10.5)
    
    # 收集所有需要翻译的段落
    paragraphs_to_translate = [p for p in doc.paragraphs if p.text.strip()]
    total_paragraphs = len(paragraphs_to_translate)
    
    # 使用线程池处理段落翻译
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = []
        for i, paragraph in enumerate(paragraphs_to_translate):
            future = executor.submit(process_paragraph, paragraph)
            futures.append((i, future))
        
        for i, future in enumerate(concurrent.futures.as_completed([f for _, f in futures])):
            progress_bar.progress((i + 1) / total_paragraphs)
    
    # 翻译表格
    if doc.tables:
        st.info("开始翻译表格...")
        table_progress = st.progress(0)
        total_tables = len(doc.tables)
        
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = [executor.submit(process_table, table, target_lang) 
                      for table in doc.tables]
            for i, _ in enumerate(concurrent.futures.as_completed(futures)):
                table_progress.progress((i + 1) / total_tables)
        
        table_progress.empty()
    
    progress_bar.empty()
    st.success("翻译完成!")
    return doc

def save_word_document(doc, output_path):
    """保存Word文档"""
    doc.save(output_path)

def main():
    global DOUBAO_API_KEY, ZHIPU_API_KEY
    
    # 添加API密钥输入区域
    with st.sidebar:
        st.header("API 配置")
        DOUBAO_API_KEY = st.text_input("输入豆包API Token", type="password")
        ZHIPU_API_KEY = st.text_input("输入智谱API Token", type="password")
        
        if st.button("验证Token"):
            if check_tokens(DOUBAO_API_KEY, ZHIPU_API_KEY):
                st.success("Token格式验证通过!")
            else:
                st.error("Token格式不正确，请检查!")
    
    # 检查是否已输入API密钥
    if not DOUBAO_API_KEY or not ZHIPU_API_KEY:
        st.warning("请先在左侧边栏输入豆包和智谱的API Token以使用服务")
        return
    
    # 文件上传区域
    def enforce_filename_restriction(filename, allowed_extensions):
        file_extension = os.path.splitext(filename)[1].lower()
        allowed_extensions = [ext.lower() for ext in allowed_extensions]
        if file_extension not in allowed_extensions:
            raise ValueError(f"Invalid file extension: {file_extension}. Allowed: {allowed_extensions}")
        return True

    # 允许上传的文件扩展名，支持大写和小写
    allowed_extensions = ['.pdf', '.PDF']

    # 创建文件上传组件
    uploaded_file = st.file_uploader("请上传一个 PDF 文件")

    if uploaded_file is not None:
        try:
            # 验证文件扩展名
            enforce_filename_restriction(uploaded_file.name, allowed_extensions)
            st.write("你已成功上传文件: ", uploaded_file.name)
        except ValueError as e:
            st.error(e)
    
    # 语言选择
    col1, col2 = st.columns(2)
    with col1:
        source_lang = st.selectbox("源语言", ["en", "zh"], index=0)
    with col2:
        target_lang = st.selectbox("目标语言", ["zh", "en"], index=0)
    
    if uploaded_file is not None:
        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            tmp_pdf.write(uploaded_file.getvalue())
            pdf_path = tmp_pdf.name
        
        # 转换为Word
        word_path = pdf_path.replace(".pdf", ".docx")
        pdf_to_word(pdf_path, word_path)
        
        # 提取文本
        doc, paragraphs, tables = extract_text_from_word(word_path)
        
        # 显示原始文本预览
        with st.expander("原始文本预览"):
            st.write("\n\n".join(paragraphs[:10]))
        
        # 翻译按钮
        if st.button("开始翻译"):
            if not check_tokens(DOUBAO_API_KEY, ZHIPU_API_KEY):
                st.error("API Token无效，请检查!")
                return
                
            translated_doc = translate_word_document(doc, paragraphs, source_lang, target_lang)
            
            # 保存翻译后的文档
            translated_word_path = word_path.replace(".docx", f"_translated_{target_lang}.docx")
            save_word_document(translated_doc, translated_word_path)
            
            # 提供下载
            with open(translated_word_path, "rb") as f:
                bytes_data = f.read()
            
            st.download_button(
                label="下载翻译后的Word文档",
                data=BytesIO(bytes_data),
                file_name=f'{uploaded_file.name}_zh.docx',
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # 清理临时文件
            os.unlink(pdf_path)
            os.unlink(word_path)
            os.unlink(translated_word_path)

if __name__ == "__main__":
    main()
